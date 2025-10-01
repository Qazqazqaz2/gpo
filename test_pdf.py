from docx import Document
import re
import os
import subprocess

def replace_text(paragraph, key, value):
    """
    Replaces text within a paragraph's runs
    """
    if key in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if key in inline[i].text:
                text = inline[i].text.replace(key, value)
                inline[i].text = text

def replace_underline(paragraph, key, value):
    """
    Replaces underline characters with the specified value
    """
    text = paragraph.text
    if '_____' in text or '___' in text:
        # Заменяем последовательности подчеркиваний
        new_text = re.sub(r'_+', lambda x: value if x.group() in ['_____', '___'] else x.group(), text)
        if new_text != text:
            paragraph.text = new_text

def fill_docx_template(template_path, data, output_path=None):
    """
    Fills a DOCX template with data and returns the Document object
    If output_path is provided, saves the document to that path
    """
    # Open the document template
    doc = Document(template_path)
    
    # Process paragraphs in the document
    for paragraph in doc.paragraphs:
        # Replace markers in square brackets
        for key, value in data.items():
            if f'[{key}]' in paragraph.text:
                replace_text(paragraph, f'[{key}]', value)
            elif key in paragraph.text:
                replace_text(paragraph, key, value)
        
        # Replace underlines
        for key, value in data.items():
            replace_underline(paragraph, key, value)
    
    # Process tables if present
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if f'[{key}]' in paragraph.text:
                            replace_text(paragraph, f'[{key}]', value)
                        elif key in paragraph.text:
                            replace_text(paragraph, key, value)
                        replace_underline(paragraph, key, value)
    
    # Save the document if output_path is provided
    if output_path:
        doc.save(output_path)
        print(f'Документ успешно заполнен и сохранен как {output_path}')
    
    return doc

def convert_to_pdf(docx_path_or_buffer, pdf_output_path=None):
    """
    Converts a DOCX file to PDF using available methods
    Returns the path to the PDF or the PDF content as bytes
    """
    # If input is a buffer, save to temp file first
    is_buffer = not isinstance(docx_path_or_buffer, str)
    temp_docx = None
    temp_dir = None
    temp_pdf = None
    pdf_content = None
    
    try:
        if is_buffer:
            import tempfile
            # Save buffer to temporary file
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            docx_path_or_buffer.seek(0)
            temp_docx.write(docx_path_or_buffer.read())
            temp_docx.close()
            docx_path = temp_docx.name
        else:
            docx_path = docx_path_or_buffer
        
        # Method 1: Use docx2pdf library
        try:
            from docx2pdf import convert
            if pdf_output_path:
                convert(docx_path, pdf_output_path)
                return pdf_output_path
            else:
                # If no output path, use temp file
                import tempfile
                temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                temp_pdf.close()
                convert(docx_path, temp_pdf.name)
                
                # Read PDF content
                with open(temp_pdf.name, 'rb') as f:
                    pdf_content = f.read()
                return pdf_content
        except ImportError:
            pass
        
        # Method 2: Use Word via COM (Windows only)
        try:
            import win32com.client
            word = win32com.client.Dispatch('Word.Application')
            doc_path = os.path.abspath(docx_path)
            
            if pdf_output_path:
                pdf_path = os.path.abspath(pdf_output_path)
            else:
                # Use a temporary file
                import tempfile
                temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                temp_pdf.close()
                pdf_path = os.path.abspath(temp_pdf.name)
            
            word_doc = word.Documents.Open(doc_path)
            word_doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
            word_doc.Close()
            word.Quit()
            
            if pdf_output_path:
                return pdf_output_path
            else:
                # Read PDF content
                with open(pdf_path, 'rb') as f:
                    pdf_content = f.read()
                return pdf_content
        except Exception:
            pass
        
        # Method 3: Use LibreOffice
        try:
            import tempfile
            libre_office = 'soffice'
            if os.name == 'nt':  # Windows
                possible_paths = [
                    r'C:\Program Files\LibreOffice\program\soffice.exe',
                    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                ]
                for path in possible_paths:
                    if os.path.exists(path):
                        libre_office = path
                        break
            
            temp_dir = tempfile.mkdtemp()
            subprocess.run([
                libre_office,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', temp_dir,
                docx_path
            ], check=True)
            
            # Get the output PDF filename
            pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
            pdf_path = os.path.join(temp_dir, pdf_filename)
            
            if pdf_output_path:
                # Copy file to output path
                import shutil
                shutil.copy2(pdf_path, pdf_output_path)
                return pdf_output_path
            else:
                # Read PDF content
                with open(pdf_path, 'rb') as f:
                    pdf_content = f.read()
                return pdf_content
        except Exception as e:
            raise Exception(f'Не удалось конвертировать DOCX в PDF: {str(e)}')
    
    finally:
        # Clean up temporary files
        if temp_docx and os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)
        if temp_pdf and os.path.exists(temp_pdf.name):
            os.unlink(temp_pdf.name)
        if temp_dir and os.path.exists(temp_dir):
            import shutil
            shutil.rmtree(temp_dir)
    
    # If we got here, all conversion methods failed
    raise Exception('Не удалось конвертировать DOCX в PDF. Установите одну из библиотек: docx2pdf, pywin32 (для Windows) или LibreOffice')

def process_template(template_path, data, output_docx=None, output_pdf=None):
    """
    Complete process: Fill template and convert to PDF
    Returns either the path to the output file or the file content as bytes
    """
    # Fill the template
    if output_docx:
        doc = fill_docx_template(template_path, data, output_docx)
        docx_path = output_docx
    else:
        import io
        doc_buffer = io.BytesIO()
        doc = fill_docx_template(template_path, data)
        doc.save(doc_buffer)
        docx_path = doc_buffer
    
    # Convert to PDF if requested
    if output_pdf or output_docx is None:  # Always convert if no specific output requested
        return convert_to_pdf(docx_path, output_pdf)
    else:
        return docx_path

# Example usage when script is run directly
if __name__ == "__main__":
    # Sample data
    data = {
        'ГРУППА0': '722-1',
        'ФИОСТУДЕНТА': 'Ермолов Артем',
        'НОМЕРСТУДЕНТА': '+79937412239',
        'МАИЛ': 'vip.bam444@mail.ru',
        'ОРГАНИЗАЦИЯ': 'ООО Невтегаз << Сибирь >>',
        'АДРЕС': 'ул.Федора Боброва д.12',
        'РУКОВОДИТЕЛЬ': 'Д.И. Новохрестова',
        'ДАТА': '05.05.2025'
    }
    
    # Use the template from the current directory
    template_path = 'ShABLON_732_grupp_Zayavlenie_na_prokhozhdenie_praktiki-1.docx'
    
    # Process the template and save both DOCX and PDF
    output_docx = 'Заявление_на_практику_заполненное.docx'
    output_pdf = 'Заявление_на_практику_заполненное.pdf'
    
    try:
        result = process_template(template_path, data, output_docx, output_pdf)
        print(f'Документы успешно созданы: {output_docx} и {output_pdf}')
    except Exception as e:
        print(f'Ошибка при обработке шаблона: {str(e)}')