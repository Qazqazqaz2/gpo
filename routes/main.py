from flask import Blueprint, render_template, redirect, url_for, request, flash, current_app, send_file, abort
from flask_login import login_required, current_user
from models import (
    User,
    Role,
    Student,
    Group,
    PracticeType,
    Contract,
    Organization,
    AskForm,
    Status,
    ConsultantGroup,
    PracticeDiary,
)
from extensions import db
from werkzeug.security import generate_password_hash, check_password_hash
import os
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from pypdf import PdfWriter, PdfReader
import tempfile
import shutil
from test_pdf import replace_text, replace_underline, process_template, convert_to_pdf
from docx import Document
import re
from datetime import datetime
from flask import session
import logging
from sqlalchemy import func

main = Blueprint('main', __name__)


NAME_TOKEN_SANITIZER = re.compile(r'[^A-Za-zА-Яа-яЁё-]')


def parse_user_full_name(username: str):
    """Разбить имя пользователя на ФИО и отбросить технические суффиксы."""
    if not username:
        return ("", "", "")
    
    cleaned = username.replace('.', ' ').replace('_', ' ')
    parts = []
    
    for raw_part in cleaned.split():
        normalized = NAME_TOKEN_SANITIZER.sub('', raw_part.strip())
        if normalized:
            parts.append(normalized.title())
    
    if not parts:
        return ("", "", "")
    
    surname = parts[0] if len(parts) > 0 else ""
    name = parts[1] if len(parts) > 1 else ""
    patronymic = ""
    
    if len(parts) > 2:
        patronymic = " ".join(parts[2:])
    
    return surname, name, patronymic


def _pick_preferred_student(students):
    """Выбрать наиболее корректную запись студента среди дублей."""
    def score(student):
        value = 0
        if student.patronymic:
            value += 2
            if not student.patronymic.isdigit():
                value += 3
        if student.name and not student.name.islower():
            value += 1
        if student.surname and not student.surname.islower():
            value += 1
        return value
    
    return max(students, key=lambda s: (score(s), -s.id))


def _merge_duplicate_students(primary, duplicates):
    """Переназначить заявки и удалить дубль студента."""
    if not duplicates:
        return
    
    changed = False
    for duplicate in duplicates:
        if duplicate.id == primary.id:
            continue
        ask_forms = AskForm.query.filter_by(student_id=duplicate.id).all()
        for ask_form in ask_forms:
            ask_form.student_id = primary.id
        db.session.delete(duplicate)
        changed = True
    
    if changed:
        db.session.commit()


def find_student_for_user(user):
    """Попробовать найти запись студента, связанную с пользователем."""
    if not user or not user.username:
        return None
    
    username = user.username.strip()
    username_lower = username.lower()
    
    # Попытка сопоставить по номеру студенческого
    if username_lower:
        student = Student.query.filter(func.lower(Student.student_id) == username_lower).first()
        if student:
            return student
    
    surname, name, patronymic = parse_user_full_name(username)
    surname_lower = surname.lower() if surname else None
    name_lower = name.lower() if name else None
    patronymic_lower = patronymic.lower() if patronymic else None
    
    def pick_with_filters(*filters):
        if not filters:
            return None
        students = Student.query.filter(*filters).all()
        if not students:
            return None
        if len(students) == 1:
            return students[0]
        primary = _pick_preferred_student(students)
        duplicates = [s for s in students if s.id != primary.id]
        _merge_duplicate_students(primary, duplicates)
        return primary
    
    if surname_lower and name_lower:
        base_filters = [
            func.lower(Student.surname) == surname_lower,
            func.lower(Student.name) == name_lower
        ]
        if patronymic_lower:
            with_patronymic = pick_with_filters(*(base_filters + [func.lower(Student.patronymic) == patronymic_lower]))
            if with_patronymic:
                return with_patronymic
        student = pick_with_filters(*base_filters)
        if student:
            return student
    
    if surname_lower:
        student = Student.query.filter(func.lower(Student.surname) == surname_lower).first()
        if student:
            return student
    
    if name_lower:
        student = Student.query.filter(func.lower(Student.name) == name_lower).first()
        if student:
            return student
    
    return None


def safe_value(value, default=""):
    return value if value else default


def fill_multiline_after(paragraphs, start_index, text):
    """Заполнить последовательность параграфов многострочным текстом, начиная с индекса."""
    if not text:
        return
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return
    idx = start_index
    for line in lines:
        while idx < len(paragraphs) and paragraphs[idx].text.strip().isdigit():
            idx += 1
        if idx < len(paragraphs):
            paragraphs[idx].text = line
            idx += 1
        else:
            break


def build_practice_diary_document(ask_form, diary):
    """Сформировать документ дневника практики на основе шаблона."""
    template_path = os.path.join(current_app.root_path, 'ShABLON_Dnevnik_praktiki_A4.docx')
    if not os.path.exists(template_path):
        raise FileNotFoundError("Шаблон дневника 'ShABLON_Dnevnik_praktiki_A4.docx' не найден.")
    
    doc = Document(template_path)
    paragraphs = doc.paragraphs
    
    student = diary.student or ask_form.student
    student_full_name = safe_value(student.full_name if student else None)
    faculty = safe_value(diary.faculty, 'Факультет безопасности')
    course = safe_value(diary.course)
    group_name = safe_value(diary.group_name, ask_form.group.name if ask_form.group else '')
    practice_place = safe_value(diary.practice_place, ask_form.contract.organization.name if ask_form.contract and ask_form.contract.organization else '')
    practice_period = safe_value(diary.practice_period)
    if not practice_period and ask_form.contract and ask_form.contract.date_start and ask_form.contract.date_end:
        practice_period = f"{ask_form.contract.date_start.strftime('%d.%m.%Y')} - {ask_form.contract.date_end.strftime('%d.%m.%Y')}"
    
    practice_type_name = safe_value(ask_form.practice_type.name if ask_form.practice_type else '')
    practice_view = safe_value(diary.assignment_theme, practice_type_name)
    
    def update_paragraph_contains(substr, new_text):
        for paragraph in paragraphs:
            if substr in paragraph.text:
                paragraph.text = new_text
                return paragraph
        return None
    
    update_paragraph_contains('ТИП практике', f"по {practice_type_name} практике: практика {practice_view}")
    
    student_signature_text = safe_value(diary.student_signature)
    if diary.student_signed_at:
        student_signature_text = f"{student_signature_text} ({diary.student_signed_at.strftime('%d.%m.%Y %H:%M')})" if student_signature_text else diary.student_signed_at.strftime('%d.%m.%Y %H:%M')
    update_paragraph_contains('С инструкцией ознакомлен', f"С инструкцией ознакомлен: {student_signature_text}")
    update_paragraph_contains('Подпись обучающегося', f"Подпись обучающегося: {student_signature_text}")
    
    update_paragraph_contains('Фамилия, имя, отчество обучающегося', f"1.\tФамилия, имя, отчество обучающегося: {student_full_name}")
    update_paragraph_contains('____________________________________________________', f"Номер студенческого билета: {safe_value(student.student_id if student else None)}")
    update_paragraph_contains('Факультет', f"2.\tФакультет: {faculty}")
    update_paragraph_contains('3. Курс', f"3. Курс {course}    4. Группа {group_name}")
    update_paragraph_contains('5. Место практики', f"5. Место практики: {practice_place}")
    update_paragraph_contains('6. Срок практики', f"6. Срок практики: {practice_period}")
    update_paragraph_contains('Рабочий график (план) проведения практики', f"Рабочий график (план) проведения практики: {safe_value(diary.work_plan)}")
    
    update_paragraph_contains('1. Тема практики', f"1. Тема практики: {safe_value(diary.assignment_theme)}")
    update_paragraph_contains('2. Цель практики', f"2. Цель практики: {safe_value(diary.assignment_goal)}")
    update_paragraph_contains('3. Задачи практики', f"3. Задачи практики: {safe_value(diary.assignment_tasks)}")
    
    content_paragraph = update_paragraph_contains('3. Содержание работ практики', '3. Содержание работ практики')
    if content_paragraph:
        content_index = paragraphs.index(content_paragraph) + 1
        fill_multiline_after(paragraphs, content_index, safe_value(diary.daily_entries))
    
    instruction_paragraph = update_paragraph_contains('4. Отметки о прохождении инструктажа', '4. Отметки о прохождении инструктажа')
    if instruction_paragraph:
        instr_index = paragraphs.index(instruction_paragraph) + 1
        fill_multiline_after(paragraphs, instr_index, safe_value(diary.instruction_notes))
    
    update_paragraph_contains('Заключение о работе обучающегося', f"а) Заключение о работе обучающегося в период практики: {safe_value(diary.evaluation_note)}")
    update_paragraph_contains('поощрения и взыскания', f"б) поощрения и взыскания (по приказам): {safe_value(diary.evaluation_rewards)}")
    update_paragraph_contains('Оценка за практику:', f"Оценка за практику: {safe_value(diary.evaluation_grade)}")
    
    update_paragraph_contains('Заключение руководителя практики от Университета', f"6. Заключение руководителя практики от Университета: {safe_value(diary.university_conclusion)}")
    update_paragraph_contains('Оценка за практику:\t', f"Оценка за практику: {safe_value(diary.university_grade)}")
    
    consultant_signature_text = safe_value(diary.consultant_signature, safe_value(ask_form.consultant_user.username if ask_form.consultant_user else ''))
    if diary.consultant_signed_at and consultant_signature_text:
        consultant_signature_text = f"{consultant_signature_text} ({diary.consultant_signed_at.strftime('%d.%m.%Y %H:%M')})"
    
    practice_leader_signature_text = safe_value(diary.practice_leader_signature, safe_value(ask_form.practice_leader_user.username if ask_form.practice_leader_user else ''))
    if diary.practice_leader_signed_at and practice_leader_signature_text:
        practice_leader_signature_text = f"{practice_leader_signature_text} ({diary.practice_leader_signed_at.strftime('%d.%m.%Y %H:%M')})"
    
    update_paragraph_contains('Руководитель практики от Университета', f"Руководитель практики от Университета: {practice_leader_signature_text}")
    
    # Добавим блок с подписью консультанта, если он есть
    if consultant_signature_text:
        doc.add_paragraph(f"Руководитель практики от профильной организации: {consultant_signature_text}")
    
    # Дата подписи руководителя
    leader_date = diary.practice_leader_signed_at or diary.consultant_signed_at or diary.student_signed_at
    if leader_date:
        update_paragraph_contains('«____» _____________  20__г.', f"«{leader_date.strftime('%d')}» {leader_date.strftime('%m')} {leader_date.strftime('%Y')} г.")
    else:
        update_paragraph_contains('«____» _____________  20__г.', f"«{datetime.now().strftime('%d')}» {datetime.now().strftime('%m')} {datetime.now().strftime('%Y')} г.")
    
    return doc


def prepare_practice_diary_docx(ask_form):
    diary = ask_form.diary
    if not diary:
        raise ValueError('Дневник ещё не заполнен.')
    
    doc = build_practice_diary_document(ask_form, diary)
    temp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(temp_dir, f"practice_diary_{ask_form.id}.docx")
    doc.save(docx_path)
    return docx_path, temp_dir

@main.route('/')
def index():
    current_app.logger.info(f"INDEX ROUTE: Access attempt from IP={request.remote_addr}")
    
    if current_user.is_authenticated:
        current_app.logger.info(f"INDEX ROUTE: Authenticated user '{current_user.username}' (ID: {current_user.id})")
        current_app.logger.info(f"INDEX ROUTE: User roles: {current_user.roles}")
        current_app.logger.info(f"INDEX ROUTE: Is teacher: {current_user.is_teacher}")
        current_app.logger.info(f"INDEX ROUTE: Is student: {current_user.is_student}")
        
        # Redirect to respective dashboard based on user role
        if current_user.is_teacher:
            current_app.logger.info(f"INDEX ROUTE: Redirecting teacher '{current_user.username}' to teacher dashboard")
            return redirect(url_for('main.teacher_dashboard'))
        elif current_user.is_consultant:
            current_app.logger.info(f"INDEX ROUTE: Redirecting consultant '{current_user.username}' to consultant dashboard")
            return redirect(url_for('main.consultant_dashboard'))
        else:
            current_app.logger.info(f"INDEX ROUTE: Redirecting student '{current_user.username}' to student dashboard")
            return redirect(url_for('main.student_dashboard'))
    
    current_app.logger.info("INDEX ROUTE: User not authenticated, showing login page")
    return render_template('login.html')

@main.route('/profile')
@login_required
def profile():
    student_info = None
    if 'студент' in current_user.roles:
        # Get student record for current user
        student = find_student_for_user(current_user)
        if student and student.group:
            student_info = {
                'group': student.group.name,
                'surname': student.surname,
                'name': student.name,
                'patronymic': student.patronymic
            }
    
    return render_template('profile.html', user=current_user, student_info=student_info)

@main.route('/student/dashboard')
@login_required
def student_dashboard():
    current_app.logger.info(f"STUDENT DASHBOARD: Access attempt by user '{current_user.username}' (ID: {current_user.id})")
    current_app.logger.info(f"STUDENT DASHBOARD: User roles: {current_user.roles}")
    current_app.logger.info(f"STUDENT DASHBOARD: IP={request.remote_addr}")
    
    # Check if user is a student
    if 'студент' not in current_user.roles:
        current_app.logger.warning(f"STUDENT DASHBOARD ACCESS DENIED: User '{current_user.username}' is not a student. Roles: {current_user.roles}")
        flash('У вас нет доступа к этой странице', 'danger')
        return redirect(url_for('main.index'))
    
    current_app.logger.info(f"STUDENT DASHBOARD: Access granted to student '{current_user.username}'")
    
    # Get all forms for the current student
    ask_forms = []
    student = find_student_for_user(current_user)
    
    if student:
        current_app.logger.info(f"STUDENT DASHBOARD: Found student record for '{current_user.username}' (ID: {student.id})")
        ask_forms = AskForm.query.filter(AskForm.student_id == student.id).all()
        current_app.logger.info(f"STUDENT DASHBOARD: Found {len(ask_forms)} forms for student")
    else:
        current_app.logger.warning(f"STUDENT DASHBOARD: No student record found for user '{current_user.username}'")
    
    current_app.logger.info(f"STUDENT DASHBOARD: Rendering dashboard for '{current_user.username}' with {len(ask_forms)} forms")
    return render_template('student_dashboard.html', ask_forms=ask_forms)


@main.route('/student/diaries')
@login_required
def student_diaries():
    current_app.logger.info(f"STUDENT DIARIES: Access attempt by user '{current_user.username}' (ID: {current_user.id})")
    
    if not current_user.is_student:
        flash('У вас нет доступа к этой странице', 'danger')
        return redirect(url_for('main.index'))
    
    ask_forms = AskForm.query.filter_by(responsible_user_id=current_user.id).order_by(AskForm.created_at.desc()).all()
    current_app.logger.info(f"STUDENT DIARIES: Found {len(ask_forms)} forms for diaries")
    
    return render_template('student_diaries.html', ask_forms=ask_forms)

@main.route('/teacher/dashboard')
@login_required
def teacher_dashboard():
    current_app.logger.info(f"TEACHER DASHBOARD: Access attempt by user '{current_user.username}' (ID: {current_user.id})")
    current_app.logger.info(f"TEACHER DASHBOARD: User roles: {current_user.roles}")
    current_app.logger.info(f"TEACHER DASHBOARD: IP={request.remote_addr}")
    
    # Check if user is a teacher
    if not current_user.is_teacher:
        current_app.logger.warning(f"TEACHER DASHBOARD ACCESS DENIED: User '{current_user.username}' is not a teacher. Roles: {current_user.roles}")
        flash('У вас нет доступа к этой странице', 'danger')
        return redirect(url_for('main.index'))
    
    current_app.logger.info(f"TEACHER DASHBOARD: Access granted to teacher '{current_user.username}'")
    
    # Get all groups for the dropdown
    groups = Group.query.order_by(Group.name).all()
    current_app.logger.info(f"TEACHER DASHBOARD: Found {len(groups)} groups")
    
    current_app.logger.info(f"TEACHER DASHBOARD: Rendering dashboard for teacher '{current_user.username}' with {len(groups)} groups")
    return render_template('teacher_dashboard.html', groups=groups)


@main.route('/consultant/dashboard')
@login_required
def consultant_dashboard():
    current_app.logger.info(f"CONSULTANT DASHBOARD: Access attempt by user '{current_user.username}' (ID: {current_user.id})")
    current_app.logger.info(f"CONSULTANT DASHBOARD: User roles: {current_user.roles}")
    current_app.logger.info(f"CONSULTANT DASHBOARD: IP={request.remote_addr}")
    
    if not current_user.is_consultant:
        current_app.logger.warning(f"CONSULTANT DASHBOARD ACCESS DENIED: User '{current_user.username}' is not a consultant. Roles: {current_user.roles}")
        flash('У вас нет доступа к этой странице', 'danger')
        return redirect(url_for('main.index'))
    
    assignments = (
        ConsultantGroup.query
        .filter_by(consultant_id=current_user.id)
        .join(Group, ConsultantGroup.group_id == Group.id)
        .order_by(Group.name)
        .all()
    )
    groups = [assignment.group for assignment in assignments if assignment.group]
    assigned_group_ids = [assignment.group_id for assignment in assignments]
    current_app.logger.info(f"CONSULTANT DASHBOARD: Assigned groups count for '{current_user.username}': {len(groups)}")
    
    selected_group_id = request.args.get('group_id', type=int)
    if not selected_group_id or selected_group_id not in assigned_group_ids:
        selected_group_id = assigned_group_ids[0] if assigned_group_ids else None
    
    forms = []
    selected_group = None
    if selected_group_id:
        forms = (
            AskForm.query
            .filter(
                AskForm.consultant_leader_id == current_user.id,
                AskForm.group_id == selected_group_id
            )
            .order_by(AskForm.created_at.desc())
            .all()
        )
        selected_group = Group.query.get(selected_group_id)
        current_app.logger.info(f"CONSULTANT DASHBOARD: Selected group ID={selected_group_id}, forms found={len(forms)}")
    
    available_groups_query = Group.query
    if assigned_group_ids:
        available_groups_query = available_groups_query.filter(~Group.id.in_(assigned_group_ids))
    available_groups = available_groups_query.order_by(Group.name).all()
    current_app.logger.info(f"CONSULTANT DASHBOARD: Available groups to claim: {len(available_groups)}")
    
    return render_template(
        'consultant_dashboard.html',
        groups=groups,
        forms=forms,
        selected_group=selected_group,
        selected_group_id=selected_group_id,
        available_groups=available_groups
    )


@main.route('/consultant/groups/<int:group_id>/claim', methods=['POST'])
@login_required
def consultant_claim_group(group_id):
    if not current_user.is_consultant:
        flash('У вас нет доступа к этой функции', 'danger')
        return redirect(url_for('main.index'))
    
    group = Group.query.get_or_404(group_id)
    
    existing_assignment = ConsultantGroup.query.filter_by(group_id=group_id).first()
    if existing_assignment:
        if existing_assignment.consultant_id == current_user.id:
            flash('Эта группа уже закреплена за вами.', 'info')
        else:
            flash('Группа уже закреплена за другим преподавателем-консультантом.', 'warning')
        return redirect(url_for('main.consultant_dashboard'))
    
    assignment = ConsultantGroup(consultant_id=current_user.id, group_id=group_id)
    db.session.add(assignment)
    db.session.commit()
    
    flash(f'Группа {group.name} закреплена за вами.', 'success')
    return redirect(url_for('main.consultant_dashboard', group_id=group_id))


@main.route('/consultant/sign/<int:form_id>/<string:document>', methods=['POST'])
@login_required
def consultant_sign(form_id, document):
    if not current_user.is_consultant:
        flash('У вас нет доступа к этой функции', 'danger')
        return redirect(url_for('main.index'))
    
    ask_form = AskForm.query.get_or_404(form_id)
    
    if ask_form.consultant_leader_id != current_user.id:
        flash('Вы не являетесь консультантом для этой заявки', 'danger')
        return redirect(url_for('main.consultant_dashboard'))
    
    signature_text = f"Подписано {current_user.username}"
    now = datetime.utcnow()
    
    if document == 'contract':
        ask_form.consultant_contract_signature = signature_text
        ask_form.consultant_contract_signed_at = now
        message = 'Договор подписан.'
    elif document == 'application':
        ask_form.consultant_application_signature = signature_text
        ask_form.consultant_application_signed_at = now
        message = 'Заявление подписано.'
    else:
        abort(400)
    
    ask_form.updated_at = now
    db.session.commit()
    
    flash(message, 'success')
    return redirect(request.referrer or url_for('main.consultant_dashboard'))

@main.route('/teacher/students/<int:group_id>')
@login_required
def students_by_group(group_id):
    # Check if user is a teacher
    if not (current_user.is_teacher or current_user.is_consultant):
        flash('У вас нет доступа к этой странице', 'danger')
        return redirect(url_for('main.index'))
    
    if current_user.is_consultant:
        assignment = ConsultantGroup.query.filter_by(consultant_id=current_user.id, group_id=group_id).first()
        if not assignment:
            flash('Эта группа не закреплена за вами.', 'danger')
            return redirect(url_for('main.consultant_dashboard'))
    
    # Get all students in the group
    students = Student.query.filter_by(group_id=group_id).all()
    group = Group.query.get_or_404(group_id)
    
    return render_template('student_list.html', students=students, group=group)

@main.route('/practice-form', methods=['GET', 'POST'])
@login_required
def practice_form():
    current_app.logger.info(f"PRACTICE FORM: Access attempt by user '{current_user.username}' (ID: {current_user.id})")
    current_app.logger.info(f"PRACTICE FORM: User roles: {current_user.roles}")
    current_app.logger.info(f"PRACTICE FORM: IP={request.remote_addr}")
    current_app.logger.info(f"PRACTICE FORM: Method={request.method}")
    
    # Check if user is a student
    if 'студент' not in current_user.roles:
        current_app.logger.warning(f"PRACTICE FORM ACCESS DENIED: User '{current_user.username}' is not a student. Roles: {current_user.roles}")
        flash('У вас нет доступа к этой странице', 'danger')
        return redirect(url_for('main.index'))
    
    current_app.logger.info(f"PRACTICE FORM: Access granted to student '{current_user.username}'")
    
    if request.method == 'POST':
        current_app.logger.info("PRACTICE FORM: Processing POST request")
        try:
            # Get form data
            practice_type = request.form.get('practice_type')
            group = request.form.get('group')
            student = request.form.get('student')
            consultant_leader = request.form.get('consultant_leader')
            practice_leader = request.form.get('practice_leader')
            phone_number = request.form.get('phone_number')
            email = request.form.get('email')
            try:
                group_id_int = int(group) if group else None
            except (TypeError, ValueError):
                group_id_int = None
            assigned_consultant = None
            if group_id_int:
                assigned_consultant = ConsultantGroup.query.filter_by(group_id=group_id_int).first()
            if (not consultant_leader or consultant_leader.strip() == '') and assigned_consultant:
                consultant_leader = str(assigned_consultant.consultant_id)
            
            if not consultant_leader:
                flash('Выберите преподавателя-консультанта.', 'danger')
                return redirect(url_for('main.practice_form'))
            
            if not practice_leader:
                flash('Укажите руководителя практики.', 'danger')
                return redirect(url_for('main.practice_form'))
            
            current_app.logger.info(f"PRACTICE FORM: Form data received:")
            current_app.logger.info(f"  - practice_type: {practice_type}")
            current_app.logger.info(f"  - group: {group}")
            current_app.logger.info(f"  - student: {student}")
            current_app.logger.info(f"  - consultant_leader: {consultant_leader}")
            current_app.logger.info(f"  - practice_leader: {practice_leader}")
            current_app.logger.info(f"  - phone_number: {phone_number}")
            current_app.logger.info(f"  - email: {email}")
            
            # Check if using custom organization
            use_custom_org = 'use_custom_org' in request.form
            current_app.logger.info(f"PRACTICE FORM: use_custom_org: {use_custom_org}")
        except Exception as e:
            current_app.logger.error(f"PRACTICE FORM ERROR: Exception during form data processing: {str(e)}")
            current_app.logger.error(f"PRACTICE FORM ERROR: Exception type: {type(e).__name__}")
            import traceback
            current_app.logger.error(f"PRACTICE FORM ERROR: Traceback: {traceback.format_exc()}")
            flash('Произошла ошибка при обработке данных заявки. Попробуйте снова.', 'danger')
            return redirect(url_for('main.practice_form'))
        
        # Store user data in session for PDF generation
        session['phone_number'] = phone_number
        session['email'] = email
        
        # Handle contract selection or custom organization
        if use_custom_org:
            # Get custom organization data
            custom_org_name = request.form.get('custom_org_name')
            custom_org_address = request.form.get('custom_org_address')
            custom_contract_num = request.form.get('custom_contract_num')
            
            # Create a new organization record
            new_org = Organization(
                name=custom_org_name,
                address=custom_org_address
            )
            db.session.add(new_org)
            db.session.flush()  # Get ID without committing
            
            # Create a new contract record
            today = datetime.now()
            new_contract = Contract(
                contract_number=custom_contract_num or f"Временный №{today.strftime('%Y%m%d%H%M%S')}",
                organization_id=new_org.id,
                date_start=today,
                date_end=today.replace(year=today.year + 1)  # One year contract
            )
            db.session.add(new_contract)
            db.session.flush()  # Get ID without committing
            
            # Use the new contract
            contract = new_contract.id
            
            # Save custom organization info in session for PDF generation
            session['custom_organization'] = True
            session['organization_name'] = custom_org_name
            session['organization_address'] = custom_org_address
        else:
            # Using existing contract
            contract = request.form.get('contract')
            session['custom_organization'] = False
        
        # Create new form with status 1 (на рассмотрении)
        status = Status.query.filter_by(name='1').first()
        if not status:
            current_app.logger.info("PRACTICE FORM: Creating new status '1' (на рассмотрении)")
            status = Status(name='1')
            db.session.add(status)
            db.session.commit()
        else:
            current_app.logger.info(f"PRACTICE FORM: Using existing status '1' (ID: {status.id})")
        
        if not group_id_int:
            flash('Некорректная группа.', 'danger')
            return redirect(url_for('main.practice_form'))
        
        try:
            practice_type_id = int(practice_type)
            consultant_leader_id = int(consultant_leader)
            practice_leader_id = int(practice_leader)
            contract_id_value = contract if isinstance(contract, int) else int(contract)
        except (TypeError, ValueError):
            current_app.logger.error("PRACTICE FORM: Invalid numeric values received.")
            flash('Некорректные данные формы. Проверьте выбранные значения.', 'danger')
            return redirect(url_for('main.practice_form'))
        
        current_student_record = find_student_for_user(current_user)
        if not current_student_record:
            current_app.logger.warning(f"PRACTICE FORM: No student record found for user '{current_user.username}'. Creating new profile.")
            surname, name, patronymic = parse_user_full_name(current_user.username)
            current_student_record = Student(
                name=name or current_user.username,
                surname=surname or current_user.username,
                patronymic=patronymic,
                group_id=group_id_int
            )
            db.session.add(current_student_record)
            db.session.commit()
            current_app.logger.info(f"PRACTICE FORM: Created student record #{current_student_record.id} for user '{current_user.username}'")
        else:
            current_app.logger.info(f"PRACTICE FORM: Found student record: {current_student_record.id}")
        
        # Create the form
        ask_form = AskForm(
            practice_type_id=practice_type_id,
            group_id=group_id_int,
            contract_id=contract_id_value,
            responsible_user_id=current_user.id,  # Current student is responsible
            consultant_leader_id=consultant_leader_id,
            practice_leader_id=practice_leader_id,
            status_id=status.id,
            student_id=current_student_record.id  # Use the actual student record ID
        )
        
        db.session.add(ask_form)
        db.session.commit()
        
        current_app.logger.info(f"PRACTICE FORM: Form created successfully with ID: {ask_form.id}")
        current_app.logger.info(f"PRACTICE FORM: Form linked to student: {current_student_record.id} ({current_student_record.name})")
        
        try:
            flash('Заявка на практику успешно отправлена!', 'success')
            current_app.logger.info(f"PRACTICE FORM: Form submitted successfully by user '{current_user.username}'")
            
            # Generate PDF
            current_app.logger.info("PRACTICE FORM: Generating PDF...")
            pdf = generate_practice_pdf(ask_form)
            current_app.logger.info("PRACTICE FORM: PDF generated successfully")
            
            # Return PDF file
            return send_file(
                BytesIO(pdf),
                mimetype='application/pdf',
                as_attachment=True,
                download_name='practice_application.pdf'
            )
            
        except Exception as e:
            current_app.logger.error(f"PRACTICE FORM ERROR: Exception during form processing: {str(e)}")
            current_app.logger.error(f"PRACTICE FORM ERROR: Exception type: {type(e).__name__}")
            import traceback
            current_app.logger.error(f"PRACTICE FORM ERROR: Traceback: {traceback.format_exc()}")
            flash('Произошла ошибка при отправке заявки. Попробуйте снова.', 'danger')
            return redirect(url_for('main.practice_form'))
    
    # GET request - show form
    current_app.logger.info("PRACTICE FORM: Processing GET request - showing form")
    
    # Get current student data
    current_student = find_student_for_user(current_user)
    if current_student:
        current_app.logger.info(f"PRACTICE FORM: Current student: {current_student.id}")
    else:
        current_app.logger.info("PRACTICE FORM: Current student profile not found for pre-fill")
    assigned_consultant_id = None
    if current_student and current_student.group_id:
        consultant_assignment = ConsultantGroup.query.filter_by(group_id=current_student.group_id).first()
        if consultant_assignment:
            assigned_consultant_id = consultant_assignment.consultant_id
            current_app.logger.info(f"PRACTICE FORM: Found consultant assignment for group {current_student.group_id}: {assigned_consultant_id}")
    
    # Get default group (722-1)
    default_group = Group.query.filter_by(name='722-1').first()
    if not default_group:
        # Don't create here, it should be created in create_defaults.py
        current_app.logger.warning("PRACTICE FORM: Group 722-1 not found, using fallback")
        flash('Группа 722-1 не найдена в базе данных', 'warning')
        default_group = Group.query.first()  # Get any group as fallback
    else:
        current_app.logger.info(f"PRACTICE FORM: Default group found: {default_group.name}")
    
    # Get data for form dropdowns
    try:
        practice_types = PracticeType.query.all()
        groups = Group.query.all()
        contracts = Contract.query.join(Organization, Contract.organization_id == Organization.id).all()
        students = Student.query.all()
        consultant_users = User.query.join(Role, User.role_id == Role.id).filter(Role.name == 'преподаватель консультант').all()
        practice_leaders = User.query.join(Role, User.role_id == Role.id).filter(Role.name == 'преподаватель').all()
        
        current_app.logger.info(f"PRACTICE FORM: Data loaded:")
        current_app.logger.info(f"  - practice_types: {len(practice_types)}")
        current_app.logger.info(f"  - groups: {len(groups)}")
        current_app.logger.info(f"  - contracts: {len(contracts)}")
        current_app.logger.info(f"  - students: {len(students)}")
        current_app.logger.info(f"  - consultants: {len(consultant_users)}")
        current_app.logger.info(f"  - practice leaders: {len(practice_leaders)}")
        
        current_app.logger.info("PRACTICE FORM: Rendering form template")
        return render_template('practice_form.html', 
                              practice_types=practice_types,
                              groups=groups,
                              contracts=contracts,
                              students=students,
                              consultants=consultant_users,
                              practice_leaders=practice_leaders,
                              current_student=current_student,
                              default_group=default_group,
                              assigned_consultant_id=assigned_consultant_id)
                              
    except Exception as e:
        current_app.logger.error(f"PRACTICE FORM ERROR: Exception during GET request: {str(e)}")
        current_app.logger.error(f"PRACTICE FORM ERROR: Exception type: {type(e).__name__}")
        import traceback
        current_app.logger.error(f"PRACTICE FORM ERROR: Traceback: {traceback.format_exc()}")
        flash('Произошла ошибка при загрузке формы. Попробуйте снова.', 'danger')
        return redirect(url_for('main.index'))

@main.route('/view-form/<int:form_id>')
@login_required
def view_form(form_id):
    ask_form = AskForm.query.get_or_404(form_id)
    
    # Check permissions
    allowed_user_ids = {
        ask_form.responsible_user_id,
        ask_form.consultant_leader_id,
        ask_form.practice_leader_id
    }
    if (not current_user.is_teacher and not current_user.is_consultant and current_user.id not in allowed_user_ids):
        flash('У вас нет доступа к этой форме', 'danger')
        return redirect(url_for('main.index'))
    
    return render_template('view_form.html', ask_form=ask_form, diary=ask_form.diary)


@main.route('/practice-diary/<int:ask_form_id>', methods=['GET', 'POST'])
@login_required
def practice_diary(ask_form_id):
    ask_form = AskForm.query.get_or_404(ask_form_id)
    
    is_owner = current_user.is_student and ask_form.responsible_user_id == current_user.id
    is_consultant_for_form = current_user.is_consultant and current_user.id == ask_form.consultant_leader_id
    is_practice_leader_for_form = current_user.is_teacher and current_user.id == ask_form.practice_leader_id
    
    if not (is_owner or is_consultant_for_form or is_practice_leader_for_form):
        flash('У вас нет доступа к этому дневнику.', 'danger')
        return redirect(url_for('main.index'))
    
    diary = ask_form.diary
    if not diary and request.method == 'GET' and is_owner:
        diary = PracticeDiary(
            ask_form_id=ask_form.id,
            student_id=ask_form.student_id,
            group_name=ask_form.group.name if ask_form.group else None
        )
        db.session.add(diary)
        db.session.commit()
        diary = ask_form.diary
    elif not diary and request.method == 'POST' and is_owner:
        diary = PracticeDiary(
            ask_form_id=ask_form.id,
            student_id=ask_form.student_id,
            group_name=ask_form.group.name if ask_form.group else None
        )
        db.session.add(diary)
    
    can_edit = is_owner
    
    if request.method == 'POST':
        if not can_edit or not diary:
            flash('Только студент может редактировать дневник.', 'danger')
            return redirect(url_for('main.practice_diary', ask_form_id=ask_form_id))
        
        diary.faculty = request.form.get('faculty')
        diary.course = request.form.get('course')
        diary.group_name = request.form.get('group_name') or (ask_form.group.name if ask_form.group else None)
        diary.practice_place = request.form.get('practice_place')
        diary.practice_period = request.form.get('practice_period')
        diary.work_plan = request.form.get('work_plan')
        diary.assignment_theme = request.form.get('assignment_theme')
        diary.assignment_goal = request.form.get('assignment_goal')
        diary.assignment_tasks = request.form.get('assignment_tasks')
        diary.daily_entries = request.form.get('daily_entries')
        diary.instruction_notes = request.form.get('instruction_notes')
        diary.evaluation_note = request.form.get('evaluation_note')
        diary.evaluation_rewards = request.form.get('evaluation_rewards')
        diary.evaluation_grade = request.form.get('evaluation_grade')
        diary.university_conclusion = request.form.get('university_conclusion')
        diary.university_grade = request.form.get('university_grade')
        
        db.session.commit()
        flash('Дневник сохранён.', 'success')
        return redirect(url_for('main.practice_diary', ask_form_id=ask_form_id))
    
    return render_template(
        'practice_diary_form.html',
        ask_form=ask_form,
        diary=diary,
        can_edit=can_edit,
        is_consultant=is_consultant_for_form,
        is_practice_leader=is_practice_leader_for_form
    )


@main.route('/practice-diary/<int:ask_form_id>/sign/<string:role>', methods=['POST'])
@login_required
def practice_diary_sign(ask_form_id, role):
    ask_form = AskForm.query.get_or_404(ask_form_id)
    diary = ask_form.diary
    
    if not diary:
        flash('Дневник ещё не заполнен студентом.', 'warning')
        return redirect(request.referrer or url_for('main.view_form', form_id=ask_form_id))
    
    signature_text = f"Подписано {current_user.username}"
    now = datetime.utcnow()
    
    if role == 'student':
        if ask_form.responsible_user_id != current_user.id:
            flash('Вы не можете подписать этот дневник.', 'danger')
            return redirect(url_for('main.practice_diary', ask_form_id=ask_form_id))
        diary.student_signature = signature_text
        diary.student_signed_at = now
        message = 'Дневник подписан студентом.'
    elif role == 'consultant':
        if current_user.id != ask_form.consultant_leader_id or not current_user.is_consultant:
            flash('Вы не являетесь консультантом для этой заявки.', 'danger')
            return redirect(url_for('main.view_form', form_id=ask_form_id))
        diary.consultant_signature = signature_text
        diary.consultant_signed_at = now
        message = 'Дневник подписан преподавателем-консультантом.'
    elif role == 'practice_leader':
        if current_user.id != ask_form.practice_leader_id or not current_user.is_teacher:
            flash('Вы не являетесь руководителем практики для этой заявки.', 'danger')
            return redirect(url_for('main.view_form', form_id=ask_form_id))
        diary.practice_leader_signature = signature_text
        diary.practice_leader_signed_at = now
        message = 'Дневник подписан руководителем практики.'
    else:
        abort(400)
    
    db.session.commit()
    flash(message, 'success')
    return redirect(request.referrer or url_for('main.view_form', form_id=ask_form_id))


@main.route('/practice-diary/<int:ask_form_id>/download/<string:file_format>')
@login_required
def practice_diary_download(ask_form_id, file_format):
    ask_form = AskForm.query.get_or_404(ask_form_id)
    diary = ask_form.diary
    
    if not diary:
        flash('Дневник ещё не заполнен студентом.', 'warning')
        return redirect(request.referrer or url_for('main.practice_diary', ask_form_id=ask_form_id))
    
    is_owner = current_user.is_student and ask_form.responsible_user_id == current_user.id
    is_consultant_for_form = current_user.is_consultant and current_user.id == ask_form.consultant_leader_id
    is_practice_leader_for_form = current_user.is_teacher and current_user.id == ask_form.practice_leader_id
    
    if not (is_owner or is_consultant_for_form or is_practice_leader_for_form):
        flash('У вас нет доступа к этому дневнику.', 'danger')
        return redirect(url_for('main.index'))
    
    temp_dir = None
    response = None
    try:
        docx_path, temp_dir = prepare_practice_diary_docx(ask_form)
        filename_base = f"practice_diary_{ask_form.id}"
        
        if file_format == 'docx':
            with open(docx_path, 'rb') as f:
                doc_bytes = f.read()
            response = send_file(
                BytesIO(doc_bytes),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"{filename_base}.docx"
            )
        elif file_format == 'pdf':
            pdf_content = convert_to_pdf(docx_path)
            if isinstance(pdf_content, bytes):
                pdf_bytes = pdf_content
            elif isinstance(pdf_content, str):
                with open(pdf_content, 'rb') as f:
                    pdf_bytes = f.read()
            else:
                raise ValueError("Не удалось сформировать PDF-файл.")
            
            response = send_file(
                BytesIO(pdf_bytes),
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f"{filename_base}.pdf"
            )
        else:
            flash('Неподдерживаемый формат файла.', 'danger')
            return redirect(request.referrer or url_for('main.practice_diary', ask_form_id=ask_form_id))
    except Exception as e:
        current_app.logger.error(f"PRACTICE DIARY DOWNLOAD ERROR: {str(e)}")
        flash('Не удалось сформировать дневник. Обратитесь к администратору.', 'danger')
        return redirect(request.referrer or url_for('main.practice_diary', ask_form_id=ask_form_id))
    finally:
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    return response

@main.route('/update-form-status/<int:form_id>/<int:status>')
@login_required
def update_form_status(form_id, status):
    # Check if user is a teacher
    if not current_user.is_teacher:
        flash('У вас нет доступа к этой функции', 'danger')
        return redirect(url_for('main.index'))
    
    ask_form = AskForm.query.get_or_404(form_id)
    status_obj = Status.query.filter_by(name=str(status)).first()
    
    if not status_obj:
        status_obj = Status(name=str(status))
        db.session.add(status_obj)
        db.session.commit()
    
    ask_form.status = status_obj
    db.session.commit()
    
    if status == 0:
        flash('Форма отклонена. Студент должен заполнить её заново.', 'warning')
    elif status == 2:
        flash('Форма принята!', 'success')
    
    return redirect(url_for('main.view_form', form_id=form_id))

def generate_practice_pdf(ask_form):
    # Path to template DOCX
    template_path = os.path.join(current_app.root_path, 'ShABLON_732_grupp_Zayavlenie_na_prokhozhdenie_praktiki-1.docx')
    
    # Get related data
    student = Student.query.get(ask_form.student_id)
    group = Group.query.get(ask_form.group_id)
    practice_type = PracticeType.query.get(ask_form.practice_type_id)
    
    # Check if using custom organization
    use_custom_org = session.get('custom_organization', False)
    
    if use_custom_org:
        # Use the custom organization data from session
        organization_name = session.get('organization_name', '')
        organization_address = session.get('organization_address', '')
    else:
        # Use organization from the selected contract
        contract = Contract.query.get(ask_form.contract_id)
        organization = Organization.query.get(contract.organization_id)
        organization_name = organization.name
        organization_address = organization.address
    
    consultant = User.query.get(ask_form.consultant_leader_id)
    practice_leader = User.query.get(ask_form.practice_leader_id)
    
    # Get phone number and email from session
    phone_number = session.get('phone_number', '+7XXXXXXXXXX')
    email = session.get('email', 'student@example.com')
    
    # Prepare data for template filling
    full_student_name = f"{student.surname} {student.name} {student.patronymic}"
    today_date = datetime.now().strftime('%d.%m.%Y')
    
    # Data for filling the docx template
    data = {
        'ГРУППА0': group.name,
        'ФИОСТУДЕНТА': full_student_name,
        'НОМЕРСТУДЕНТА': phone_number,
        'МАИЛ': email,
        'ОРГАНИЗАЦИЯ': organization_name,
        'АДРЕС': organization_address,
        'РУКОВОДИТЕЛЬ': practice_leader.username,
        'ДАТА': today_date
    }
    
    try:
        # Process the template and generate PDF
        pdf_data = process_template(template_path, data)
        
        # If a string is returned (file path), read the file
        if isinstance(pdf_data, str):
            with open(pdf_data, 'rb') as f:
                return f.read()
        
        # Otherwise, the binary data was returned directly
        return pdf_data
        
    except Exception as e:
        current_app.logger.error(f"Error generating PDF: {str(e)}")
        raise Exception(f"Ошибка при генерации PDF: {str(e)}") 